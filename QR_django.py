#pip install python-docx
#pip install pillow
#pip install qrcode

import qrcode
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

for i in range(50): # Crear 50 documentos

    input_data = str(i) # recorre i para crear los QR con valor = i

    #Creating an instance of qrcode
    qr = qrcode.QRCode(
            version=1,
            box_size=10,
            border=5)
    qr.add_data(input_data) # Asignar valor i al QR
    qr.make(fit=True)
    img = qr.make_image(fill='black', back_color='white')
    img.save("img/"+str(i)+'.png')  # Guarda una imagen con el QR

    document = Document()
    # h = document.add_heading('BAT NUMBER', 0) # Título del documento.
    # h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(200) # Tamaño a la letra

    p = document.add_paragraph(str(i))  # Agrega texto al documento. Le agrega "i" para que se vea el número a escanear.
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER    # Centra el texto

    p.style = document.styles['Normal'] # Aplica el estilo

    p = document.add_paragraph()    # Agrega otro párrafo para agregar el QR
    run = p.add_run()
    picture = run.add_picture("img/"+str(i)+'.png', width=Inches(3.25)) # Lee la imagen guardada y la agrega al documento.
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER    # Centra la imagen en el documento

    document.save("doc/"+str(i)+".docx")    # Guarda el documento con nombre "i"
    print("Creado el doc",str(i))
